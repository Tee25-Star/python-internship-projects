from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_report():
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('Habit Tracking Application', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.bold = True
    
    doc.add_paragraph()
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Requirements Analysis',
        '2. System Architecture Design',
        '3. MVP (Minimum Viable Product) Implementation',
        '4. Testing',
        '5. Documentation',
        '6. Final Presentation'
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para_format = para.runs[0].font
        para_format.size = Pt(11)
    
    doc.add_page_break()
    
    # 1. Requirements Analysis
    doc.add_heading('1. Requirements Analysis', 1)
    
    doc.add_heading('1.1 Functional Requirements', 2)
    functional_reqs = [
        'The system shall allow users to create new habits with custom names',
        'The system shall track daily completion status for each habit',
        'The system shall calculate and display current streak for each habit',
        'The system shall calculate and display longest streak achieved for each habit',
        'The system shall display total number of completion days for each habit',
        'The system shall allow users to mark habits as complete for the current day',
        'The system shall allow users to unmark completion if marked by mistake',
        'The system shall allow users to delete habits they no longer wish to track',
        'The system shall prevent duplicate habit names',
        'The system shall persist habit data between application sessions',
        'The system shall provide a modern, intuitive graphical user interface'
    ]
    
    for req in functional_reqs:
        para = doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('1.2 Non-Functional Requirements', 2)
    non_functional_reqs = [
        'The application shall use a modern, visually appealing dark theme interface',
        'The application shall be responsive and handle multiple habits efficiently',
        'The application shall save data locally in JSON format',
        'The application shall load existing data automatically on startup',
        'The user interface shall be intuitive and require minimal learning curve',
        'The application shall provide visual feedback for user actions',
        'The application shall display statistics in a clear, organized manner'
    ]
    
    for req in non_functional_reqs:
        para = doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('1.3 User Stories', 2)
    user_stories = [
        'As a user, I want to add new habits so that I can track multiple activities',
        'As a user, I want to mark habits as complete daily so that I can track my progress',
        'As a user, I want to see my current streak so that I can stay motivated',
        'As a user, I want to see my longest streak so that I can see my best performance',
        'As a user, I want to delete habits so that I can remove ones I no longer need',
        'As a user, I want my data to be saved automatically so that I do not lose my progress'
    ]
    
    for story in user_stories:
        para = doc.add_paragraph(story, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. System Architecture Design
    doc.add_heading('2. System Architecture Design', 1)
    
    doc.add_heading('2.1 Technology Stack', 2)
    tech_stack = [
        'Programming Language: Python 3.x',
        'GUI Framework: CustomTkinter 5.2.0+',
        'Data Storage: JSON file format',
        'Date Handling: Python datetime module',
        'Type Hints: Python typing module'
    ]
    
    for tech in tech_stack:
        para = doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('2.2 System Components', 2)
    
    doc.add_paragraph('The application consists of a single main class that handles all functionality:')
    
    components = [
        'HabitTracker Class: Main application class containing all business logic and UI components',
        'Data Persistence Layer: JSON file-based storage (habits_data.json)',
        'User Interface Layer: CustomTkinter widgets for modern GUI',
        'Business Logic Layer: Streak calculation, completion tracking, and data management'
    ]
    
    for comp in components:
        para = doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_heading('2.3 Data Model', 2)
    
    doc.add_paragraph('Each habit is stored with the following structure:')
    
    data_model = doc.add_paragraph()
    data_model.add_run('habit_name: {').bold = True
    data_model.add_run('\n    "created": "YYYY-MM-DD",')
    data_model.add_run('\n    "completions": ["YYYY-MM-DD", ...],')
    data_model.add_run('\n    "current_streak": integer,')
    data_model.add_run('\n    "longest_streak": integer')
    data_model.add_run('\n}')
    
    doc.add_heading('2.4 Architecture Diagram', 2)
    doc.add_paragraph('The application follows a monolithic architecture with the following flow:')
    
    flow = [
        'User Interface (CustomTkinter)',
        '    ?',
        'HabitTracker Class (Business Logic)',
        '    ?',
        'Data Persistence (JSON File)'
    ]
    
    for item in flow:
        para = doc.add_paragraph(item)
        para_format = para.runs[0].font
        para_format.name = 'Courier New'
    
    doc.add_page_break()
    
    # 3. MVP Implementation
    doc.add_heading('3. MVP (Minimum Viable Product) Implementation', 1)
    
    doc.add_heading('3.1 Core Features Implemented', 2)
    
    doc.add_heading('3.1.1 Habit Management', 3)
    features1 = [
        'Add Habit: Users can create new habits with custom names through an input field',
        'Delete Habit: Users can remove habits they no longer want to track',
        'Duplicate Prevention: System prevents adding habits with duplicate names',
        'Input Validation: Empty habit names are rejected'
    ]
    
    for feat in features1:
        para = doc.add_paragraph(feat, style='List Bullet')
    
    doc.add_heading('3.1.2 Completion Tracking', 3)
    features2 = [
        'Mark Complete: Users can mark habits as complete for the current day',
        'Unmark Complete: Users can undo completion if marked by mistake',
        'Daily Limit: Each habit can only be marked complete once per day',
        'Visual Feedback: Button changes appearance when habit is completed'
    ]
    
    for feat in features2:
        para = doc.add_paragraph(feat, style='List Bullet')
    
    doc.add_heading('3.1.3 Statistics Display', 3)
    features3 = [
        'Current Streak: Displays consecutive days of completion ending today',
        'Longest Streak: Shows the maximum consecutive days achieved',
        'Total Completions: Displays total number of days the habit was completed',
        'Real-time Updates: Statistics update automatically when habits are marked complete'
    ]
    
    for feat in features3:
        para = doc.add_paragraph(feat, style='List Bullet')
    
    doc.add_heading('3.1.4 User Interface', 3)
    features4 = [
        'Modern Dark Theme: Professional dark mode interface',
        'Card-based Layout: Each habit displayed in an attractive card format',
        'Scrollable Interface: Handles multiple habits with scrollable frame',
        'Color-coded Statistics: Different colors for different metrics',
        'Responsive Design: Adapts to different window sizes',
        'Intuitive Controls: Clear buttons and labels for all actions'
    ]
    
    for feat in features4:
        para = doc.add_paragraph(feat, style='List Bullet')
    
    doc.add_heading('3.2 Key Methods', 2)
    
    methods = [
        ('add_habit()', 'Validates input and creates new habit entry in the system'),
        ('mark_complete()', 'Records completion for current day and updates streaks'),
        ('mark_incomplete()', 'Removes current day completion and recalculates streaks'),
        ('update_streaks()', 'Calculates current and longest streaks based on completion history'),
        ('delete_habit()', 'Removes habit from tracking system'),
        ('refresh_habits_display()', 'Updates the UI to reflect current habit data'),
        ('create_habit_card()', 'Generates visual card component for each habit'),
        ('load_habits()', 'Reads habit data from JSON file on startup'),
        ('save_habits()', 'Writes current habit data to JSON file')
    ]
    
    for method, desc in methods:
        para = doc.add_paragraph()
        para.add_run(f'{method}: ').bold = True
        para.add_run(desc)
    
    doc.add_heading('3.3 Streak Calculation Algorithm', 2)
    
    doc.add_paragraph('The streak calculation uses the following logic:')
    
    algo_steps = [
        'Sort all completion dates in ascending order',
        'For current streak: Check consecutive days backwards from today',
        'For longest streak: Iterate through sorted dates and find maximum consecutive sequence',
        'Update both values whenever completions change'
    ]
    
    for step in algo_steps:
        para = doc.add_paragraph(step, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Testing
    doc.add_heading('4. Testing', 1)
    
    doc.add_heading('4.1 Test Scenarios', 2)
    
    doc.add_heading('4.1.1 Functional Testing', 3)
    
    test_cases = [
        ('Add Habit', 'Verify new habit is created and appears in the list', 'PASS'),
        ('Add Duplicate Habit', 'Verify error message appears for duplicate names', 'PASS'),
        ('Add Empty Habit', 'Verify empty habit names are rejected', 'PASS'),
        ('Mark Complete', 'Verify habit is marked complete and streak updates', 'PASS'),
        ('Mark Incomplete', 'Verify completion can be undone', 'PASS'),
        ('Delete Habit', 'Verify habit is removed from the list', 'PASS'),
        ('Streak Calculation', 'Verify streaks are calculated correctly for various scenarios', 'PASS'),
        ('Data Persistence', 'Verify habits persist after application restart', 'PASS'),
        ('Multiple Habits', 'Verify system handles multiple habits simultaneously', 'PASS')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Test Case'
    header_cells[1].text = 'Description'
    header_cells[2].text = 'Status'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    for test_case, description, status in test_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = test_case
        row_cells[1].text = description
        row_cells[2].text = status
    
    doc.add_heading('4.2 Edge Cases Tested', 2)
    
    edge_cases = [
        'Adding habit with special characters',
        'Marking habit complete multiple times in same day',
        'Deleting habit with active streak',
        'Application behavior with no habits',
        'Streak calculation with gaps in completion dates',
        'Streak calculation with single completion',
        'Data file corruption handling'
    ]
    
    for case in edge_cases:
        para = doc.add_paragraph(case, style='List Bullet')
    
    doc.add_heading('4.3 User Acceptance Testing', 2)
    
    uat_points = [
        'Interface is intuitive and easy to navigate',
        'All buttons respond correctly to user input',
        'Statistics display accurately',
        'Visual feedback is clear and immediate',
        'Application performs smoothly with multiple habits',
        'Data persists correctly between sessions'
    ]
    
    for point in uat_points:
        para = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Documentation
    doc.add_heading('5. Documentation', 1)
    
    doc.add_heading('5.1 Code Documentation', 2)
    
    doc.add_paragraph('The application includes comprehensive docstrings for all major methods:')
    
    doc_items = [
        'Each method has a clear docstring explaining its purpose',
        'Class-level documentation describes the overall functionality',
        'Inline comments explain complex logic, particularly in streak calculation',
        'Type hints are used throughout for better code clarity'
    ]
    
    for item in doc_items:
        para = doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.2 User Documentation', 2)
    
    doc.add_heading('5.2.1 Installation Instructions', 3)
    
    install_steps = [
        'Ensure Python 3.x is installed on the system',
        'Navigate to the Habit Tracker directory',
        'Install required dependencies using: pip install -r requirements.txt',
        'Run the application using: python habit_tracker.py'
    ]
    
    for step in install_steps:
        para = doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('5.2.2 Usage Guide', 3)
    
    usage_steps = [
        'Adding a Habit: Enter the habit name in the input field and click "Add Habit"',
        'Marking Complete: Click the "Mark Complete" button on any habit card',
        'Viewing Statistics: Current streak, longest streak, and total completions are displayed on each card',
        'Deleting a Habit: Click the "Delete" button on the habit card you wish to remove',
        'Data Storage: All data is automatically saved to habits_data.json in the application directory'
    ]
    
    for step in usage_steps:
        para = doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('5.3 File Structure', 2)
    
    file_structure = [
        'habit_tracker.py - Main application file containing all code',
        'requirements.txt - Python package dependencies',
        'habits_data.json - Data storage file (created automatically)'
    ]
    
    for file in file_structure:
        para = doc.add_paragraph(file, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. Final Presentation
    doc.add_heading('6. Final Presentation', 1)
    
    doc.add_heading('6.1 Project Summary', 2)
    
    summary = doc.add_paragraph()
    summary.add_run('The Habit Tracking Application is a fully functional desktop application built with Python and CustomTkinter. ')
    summary.add_run('It provides users with an intuitive interface to track daily habits, monitor progress through streaks, ')
    summary.add_run('and maintain motivation through visual statistics. The application successfully implements all core features ')
    summary.add_run('required for effective habit tracking, including creation, completion tracking, streak calculation, and data persistence.')
    
    doc.add_heading('6.2 Key Achievements', 2)
    
    achievements = [
        'Successfully implemented a modern, visually appealing user interface using CustomTkinter',
        'Developed robust streak calculation algorithm that accurately tracks consecutive completions',
        'Implemented reliable data persistence using JSON file format',
        'Created an intuitive user experience with clear visual feedback',
        'Delivered a complete MVP with all essential features functional',
        'Ensured code quality through proper documentation and type hints'
    ]
    
    for achievement in achievements:
        para = doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_heading('6.3 Technical Highlights', 2)
    
    highlights = [
        'Object-oriented design with single cohesive class structure',
        'Efficient date handling and streak calculation algorithms',
        'Modern GUI framework providing professional appearance',
        'Automatic data persistence without user intervention',
        'Real-time UI updates based on user actions',
        'Error handling for edge cases and invalid inputs'
    ]
    
    for highlight in highlights:
        para = doc.add_paragraph(highlight, style='List Bullet')
    
    doc.add_heading('6.4 Future Enhancements', 2)
    
    enhancements = [
        'Add habit categories or tags for better organization',
        'Implement weekly and monthly statistics views',
        'Add reminder notifications for incomplete habits',
        'Export data to CSV or PDF format',
        'Add habit templates for common habits',
        'Implement dark/light theme toggle',
        'Add habit icons or emoji customization',
        'Create data visualization charts for progress tracking'
    ]
    
    for enhancement in enhancements:
        para = doc.add_paragraph(enhancement, style='List Bullet')
    
    doc.add_heading('6.5 Conclusion', 2)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('The Habit Tracking Application successfully meets all specified requirements and provides a solid foundation ')
    conclusion.add_run('for personal habit management. The application demonstrates effective use of Python GUI programming, data ')
    conclusion.add_run('persistence, and user-centered design principles. With its modern interface and reliable functionality, ')
    conclusion.add_run('the application is ready for daily use and can serve as a base for future enhancements.')
    
    # Save document
    doc.save('Habit_Tracker_Report.docx')
    print("Report generated successfully: Habit_Tracker_Report.docx")

if __name__ == "__main__":
    create_report()
