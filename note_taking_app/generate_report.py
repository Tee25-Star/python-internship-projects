from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_report():
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('Modern Note-Taking System with OCR Support', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Documentation Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.italic = True
    
    doc.add_paragraph()
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Requirements Analysis',
        '2. System Architecture Design',
        '3. MVP (Minimum Viable Product) Implementation',
        '4. Testing',
        '5. Documentation',
        '6. Final Presentation'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Requirements Analysis
    doc.add_heading('1. Requirements Analysis', 1)
    
    doc.add_heading('1.1 Project Overview', 2)
    doc.add_paragraph(
        'The Modern Note-Taking System with OCR Support is a desktop application designed to provide '
        'users with an intuitive interface for creating, managing, and organizing notes. The system '
        'incorporates Optical Character Recognition (OCR) technology to extract text from images, '
        'making it easy to digitize handwritten notes, documents, and printed materials.'
    )
    
    doc.add_heading('1.2 Functional Requirements', 2)
    func_req = [
        'FR1: The system shall allow users to create new notes with customizable titles and content.',
        'FR2: The system shall provide an interface to edit existing notes.',
        'FR3: The system shall enable users to delete notes with confirmation prompts.',
        'FR4: The system shall automatically save notes as users type (auto-save functionality).',
        'FR5: The system shall provide manual save functionality for explicit note persistence.',
        'FR6: The system shall display a list of all notes in a sidebar with previews.',
        'FR7: The system shall allow users to search and filter notes by title or content.',
        'FR8: The system shall extract text from images using OCR technology.',
        'FR9: The system shall support multiple image formats (PNG, JPG, JPEG, BMP, TIFF, GIF).',
        'FR10: The system shall allow users to add extracted OCR text to existing notes or create new notes.',
        'FR11: The system shall track creation and modification dates for each note.',
        'FR12: The system shall persist all notes to local storage in JSON format.'
    ]
    
    for req in func_req:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('1.3 Non-Functional Requirements', 2)
    non_func_req = [
        'NFR1: The system shall have a modern, visually appealing user interface with dark theme.',
        'NFR2: The system shall be responsive and handle large numbers of notes efficiently.',
        'NFR3: OCR processing shall be performed asynchronously to prevent UI freezing.',
        'NFR4: The system shall provide visual feedback during OCR processing.',
        'NFR5: The system shall be cross-platform compatible (Windows, macOS, Linux).',
        'NFR6: The system shall store data locally without requiring internet connectivity (except for initial OCR model download).',
        'NFR7: The system shall have intuitive navigation and user-friendly controls.',
        'NFR8: The system shall handle errors gracefully with informative error messages.'
    ]
    
    for req in non_func_req:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('1.4 User Stories', 2)
    user_stories = [
        'As a user, I want to create notes quickly so that I can capture my thoughts immediately.',
        'As a user, I want to search through my notes so that I can find specific information quickly.',
        'As a user, I want to extract text from images so that I can digitize printed documents.',
        'As a user, I want my notes to be saved automatically so that I do not lose my work.',
        'As a user, I want to see when notes were last modified so that I can track my activity.',
        'As a user, I want a visually appealing interface so that I enjoy using the application.'
    ]
    
    for story in user_stories:
        doc.add_paragraph(story, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. System Architecture Design
    doc.add_heading('2. System Architecture Design', 1)
    
    doc.add_heading('2.1 Architecture Overview', 2)
    doc.add_paragraph(
        'The application follows a Model-View-Controller (MVC) architectural pattern with the following components:'
    )
    
    doc.add_heading('2.2 Technology Stack', 2)
    tech_stack = [
        'Python 3.8+: Core programming language',
        'CustomTkinter: Modern GUI framework for creating the user interface',
        'EasyOCR: OCR engine for text extraction from images',
        'OpenCV: Image processing library for handling image operations',
        'Pillow (PIL): Image manipulation and processing',
        'NumPy: Numerical operations for image processing',
        'JSON: Data serialization for note storage',
        'Threading: Asynchronous processing for OCR operations'
    ]
    
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('2.3 System Components', 2)
    
    doc.add_heading('2.3.1 User Interface Layer', 3)
    doc.add_paragraph(
        'The UI layer is built using CustomTkinter and consists of:'
    )
    ui_components = [
        'Sidebar: Displays list of notes with search functionality',
        'Editor Panel: Main note editing area with title and content fields',
        'Action Buttons: Save, Delete, and OCR extraction controls',
        'Dialog Windows: OCR processing feedback and results display'
    ]
    
    for comp in ui_components:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_heading('2.3.2 Business Logic Layer', 3)
    doc.add_paragraph(
        'The NoteTakingApp class encapsulates all business logic:'
    )
    logic_components = [
        'Note Management: Create, read, update, delete operations',
        'Search and Filter: Real-time note filtering based on search queries',
        'Data Persistence: JSON-based storage and retrieval',
        'OCR Integration: Image processing and text extraction coordination'
    ]
    
    for comp in logic_components:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_heading('2.3.3 Data Layer', 3)
    doc.add_paragraph(
        'Data storage is handled through:'
    )
    data_components = [
        'Local File System: Notes stored in JSON format in a dedicated directory',
        'Note Structure: Each note contains title, content, creation date, and modification date',
        'Metadata: Timestamps and unique identifiers for each note'
    ]
    
    for comp in data_components:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_heading('2.4 Data Flow', 2)
    doc.add_paragraph(
        '1. User creates/edits a note ? UI captures input ? NoteTakingApp processes ? JSON file updated'
    )
    doc.add_paragraph(
        '2. User searches notes ? Search query processed ? Notes filtered ? UI updated with results'
    )
    doc.add_paragraph(
        '3. User uploads image for OCR ? Image loaded ? OCR thread processes ? Results displayed ? User chooses to add to note'
    )
    
    doc.add_heading('2.5 Design Patterns', 2)
    design_patterns = [
        'Singleton Pattern: Single instance of OCR reader initialized lazily',
        'Observer Pattern: UI updates automatically when notes change',
        'Threading Pattern: Asynchronous OCR processing to maintain UI responsiveness'
    ]
    
    for pattern in design_patterns:
        doc.add_paragraph(pattern, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. MVP Implementation
    doc.add_heading('3. MVP (Minimum Viable Product) Implementation', 1)
    
    doc.add_heading('3.1 Core Features Implemented', 2)
    core_features = [
        'Note Creation: Users can create new notes with titles and content',
        'Note Editing: Full text editing capabilities with auto-save',
        'Note Deletion: Safe deletion with confirmation prompts',
        'Note List Display: Sidebar showing all notes with previews',
        'Search Functionality: Real-time filtering of notes',
        'OCR Text Extraction: Extract text from images and integrate into notes',
        'Data Persistence: Automatic saving to local JSON files',
        'Date Tracking: Creation and modification timestamps'
    ]
    
    for feature in core_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.2 Key Classes and Methods', 2)
    
    doc.add_heading('3.2.1 NoteTakingApp Class', 3)
    doc.add_paragraph('Main application class containing all functionality.')
    
    methods = [
        '__init__(): Initializes the application, sets up storage, and creates UI',
        'create_ui(): Builds the complete user interface',
        'create_new_note(): Creates a new note with unique timestamp ID',
        'load_note_to_editor(): Loads selected note into editor',
        'save_note(): Manually saves current note',
        'save_note_auto(): Auto-saves note on content change',
        'delete_note(): Deletes current note with confirmation',
        'refresh_notes_list(): Updates sidebar with current notes',
        'filter_notes(): Filters notes based on search query',
        'load_notes(): Loads all notes from JSON file',
        'save_notes_to_file(): Persists notes to disk',
        'initialize_ocr(): Lazy initialization of OCR reader',
        'open_ocr_dialog(): Opens file dialog for image selection',
        'process_ocr(): Processes image and extracts text',
        'show_ocr_results(): Displays OCR results and options'
    ]
    
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_heading('3.3 User Interface Components', 2)
    
    doc.add_heading('3.3.1 Sidebar', 3)
    sidebar_features = [
        'Title header: "My Notes"',
        'New Note button: Creates new note',
        'OCR button: Opens image selection dialog',
        'Search box: Filters notes in real-time',
        'Notes list: Scrollable list with previews and dates'
    ]
    
    for feature in sidebar_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.3.2 Editor Panel', 3)
    editor_features = [
        'Title entry field: Editable note title',
        'Text editor: Multi-line text editing area',
        'Save button: Manual save trigger',
        'Delete button: Note deletion',
        'Date label: Shows last modification time'
    ]
    
    for feature in editor_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.4 OCR Implementation', 2)
    doc.add_paragraph(
        'OCR functionality is implemented using EasyOCR library:'
    )
    ocr_steps = [
        'User clicks "Extract Text from Image" button',
        'File dialog opens for image selection',
        'Loading window displays progress indicator',
        'OCR processing runs in separate thread to prevent UI freezing',
        'EasyOCR reads image and extracts text',
        'Results displayed in dialog window',
        'User can add text to current note or create new note'
    ]
    
    for step in ocr_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('3.5 Data Storage', 2)
    doc.add_paragraph(
        'Notes are stored in JSON format with the following structure:'
    )
    doc.add_paragraph(
        '{"note_id": {"title": "Note Title", "content": "Note content...", '
        '"created": "2026-01-27T19:00:00", "modified": "2026-01-27T19:30:00"}}'
    )
    doc.add_paragraph(
        'All notes are saved in a "notes" directory within the application folder, '
        'ensuring data persistence across application sessions.'
    )
    
    doc.add_page_break()
    
    # 4. Testing
    doc.add_heading('4. Testing', 1)
    
    doc.add_heading('4.1 Testing Strategy', 2)
    doc.add_paragraph(
        'Testing was performed through manual testing and functional verification of all features.'
    )
    
    doc.add_heading('4.2 Test Cases', 2)
    
    doc.add_heading('4.2.1 Note Management Tests', 3)
    test_cases = [
        'TC1: Create new note - Verify note appears in sidebar and editor is cleared',
        'TC2: Edit note title - Verify title updates in sidebar and persists after save',
        'TC3: Edit note content - Verify content saves correctly',
        'TC4: Delete note - Verify note removed from list and editor cleared',
        'TC5: Auto-save - Verify note saves automatically while typing',
        'TC6: Manual save - Verify explicit save button works correctly'
    ]
    
    for tc in test_cases:
        doc.add_paragraph(tc, style='List Bullet')
    
    doc.add_heading('4.2.2 Search Functionality Tests', 3)
    search_tests = [
        'TC7: Search by title - Verify correct notes filtered',
        'TC8: Search by content - Verify notes with matching content appear',
        'TC9: Case-insensitive search - Verify search works regardless of case',
        'TC10: Empty search - Verify all notes displayed when search cleared'
    ]
    
    for tc in search_tests:
        doc.add_paragraph(tc, style='List Bullet')
    
    doc.add_heading('4.2.3 OCR Functionality Tests', 3)
    ocr_tests = [
        'TC11: Select image file - Verify file dialog opens correctly',
        'TC12: Process PNG image - Verify text extraction works',
        'TC13: Process JPG image - Verify different format supported',
        'TC14: Add OCR text to existing note - Verify text appended correctly',
        'TC15: Create new note from OCR - Verify new note created with extracted text',
        'TC16: Cancel OCR dialog - Verify no changes when cancelled'
    ]
    
    for tc in ocr_tests:
        doc.add_paragraph(tc, style='List Bullet')
    
    doc.add_heading('4.2.4 Data Persistence Tests', 3)
    persistence_tests = [
        'TC17: Save and reload - Verify notes persist after application restart',
        'TC18: Multiple notes - Verify all notes saved correctly',
        'TC19: Note modification dates - Verify timestamps update correctly'
    ]
    
    for tc in persistence_tests:
        doc.add_paragraph(tc, style='List Bullet')
    
    doc.add_heading('4.3 Test Results', 2)
    doc.add_paragraph(
        'All test cases passed successfully. The application demonstrates:'
    )
    results = [
        'Reliable note creation and management',
        'Accurate search and filtering',
        'Successful OCR text extraction from various image formats',
        'Proper data persistence across sessions',
        'Responsive user interface without freezing during OCR processing'
    ]
    
    for result in results:
        doc.add_paragraph(result, style='List Bullet')
    
    doc.add_heading('4.4 Known Limitations', 2)
    limitations = [
        'OCR accuracy depends on image quality and clarity',
        'First OCR use requires internet connection for model download',
        'Large images may take longer to process',
        'OCR supports English language by default'
    ]
    
    for limit in limitations:
        doc.add_paragraph(limit, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Documentation
    doc.add_heading('5. Documentation', 1)
    
    doc.add_heading('5.1 Code Documentation', 2)
    doc.add_paragraph(
        'The codebase includes comprehensive inline documentation:'
    )
    doc_items = [
        'Class docstrings explaining the purpose of NoteTakingApp',
        'Method docstrings describing functionality and parameters',
        'Inline comments explaining complex logic',
        'Clear variable naming conventions'
    ]
    
    for item in doc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.2 User Documentation', 2)
    doc.add_paragraph(
        'A README.md file provides:'
    )
    readme_items = [
        'Installation instructions',
        'Usage guide for all features',
        'Troubleshooting section',
        'Requirements and dependencies list'
    ]
    
    for item in readme_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.3 API Documentation', 2)
    doc.add_paragraph(
        'Key methods and their purposes:'
    )
    
    api_table = doc.add_table(rows=1, cols=3)
    api_table.style = 'Light Grid Accent 1'
    hdr_cells = api_table.rows[0].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = 'Parameters'
    hdr_cells[2].text = 'Description'
    
    api_methods = [
        ('create_new_note', 'None', 'Creates a new note with timestamp ID'),
        ('save_note', 'None', 'Manually saves current note to file'),
        ('delete_note', 'None', 'Deletes current note after confirmation'),
        ('load_note_to_editor', 'note_id: str', 'Loads specified note into editor'),
        ('open_ocr_dialog', 'None', 'Opens file dialog for OCR image selection'),
        ('filter_notes', 'event: optional', 'Filters notes based on search query')
    ]
    
    for method, params, desc in api_methods:
        row_cells = api_table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = params
        row_cells[2].text = desc
    
    doc.add_heading('5.4 Installation Guide', 2)
    install_steps = [
        'Install Python 3.8 or higher',
        'Install dependencies: pip install -r requirements.txt',
        'Run application: python note_taking_app.py',
        'First OCR use will download language models (requires internet)'
    ]
    
    for step in install_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # 6. Final Presentation
    doc.add_heading('6. Final Presentation', 1)
    
    doc.add_heading('6.1 Project Summary', 2)
    doc.add_paragraph(
        'The Modern Note-Taking System with OCR Support successfully delivers a functional, '
        'user-friendly application that combines traditional note-taking capabilities with '
        'advanced OCR technology. The MVP includes all core features necessary for effective '
        'note management and text extraction from images.'
    )
    
    doc.add_heading('6.2 Key Achievements', 2)
    achievements = [
        'Successfully implemented a modern, visually appealing user interface',
        'Integrated OCR technology for text extraction from images',
        'Implemented auto-save functionality for seamless user experience',
        'Created efficient search and filtering system',
        'Ensured data persistence with JSON-based storage',
        'Maintained UI responsiveness through asynchronous OCR processing'
    ]
    
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_heading('6.3 Technical Highlights', 2)
    highlights = [
        'CustomTkinter for modern GUI design',
        'EasyOCR integration for accurate text extraction',
        'Threading for non-blocking OCR operations',
        'JSON-based data persistence',
        'Real-time search and filtering',
        'Auto-save mechanism'
    ]
    
    for highlight in highlights:
        doc.add_paragraph(highlight, style='List Bullet')
    
    doc.add_heading('6.4 Future Enhancements', 2)
    future_features = [
        'Support for multiple languages in OCR',
        'Note categorization and tagging system',
        'Export notes to PDF or other formats',
        'Rich text editing with formatting options',
        'Note sharing capabilities',
        'Cloud synchronization',
        'Image attachment to notes',
        'Note templates',
        'Dark/light theme toggle',
        'Keyboard shortcuts for power users'
    ]
    
    for feature in future_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('6.5 Conclusion', 2)
    doc.add_paragraph(
        'The Modern Note-Taking System with OCR Support represents a successful implementation '
        'of an MVP that addresses the core requirements of note-taking and OCR functionality. '
        'The application provides a solid foundation for future enhancements while delivering '
        'immediate value to users through its intuitive interface and powerful features.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The project demonstrates proficiency in Python GUI development, OCR integration, '
        'asynchronous programming, and software engineering best practices. All functional '
        'and non-functional requirements have been met, and the system is ready for deployment.'
    )
    
    # Save document
    filename = 'Note_Taking_System_Report.docx'
    doc.save(filename)
    print(f'Report generated successfully: {filename}')

if __name__ == '__main__':
    create_report()
