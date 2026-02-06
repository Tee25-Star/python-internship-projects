#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to generate comprehensive Word documentation for the Ticketing System project
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_code_block(doc, code):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.style = 'No Spacing'
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    para._element.get_or_add_pPr().append(shading_elm)

def create_documentation():
    """Create comprehensive project documentation"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('Service Request / Ticketing System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True
    
    doc.add_paragraph()
    date_str = datetime.now().strftime("%B %d, %Y")
    date_para = doc.add_paragraph(f'Date: {date_str}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc = [
        '1. Requirements Analysis',
        '2. System Architecture Design',
        '3. MVP (Minimum Viable Product) Implementation',
        '4. Testing',
        '5. Documentation',
        '6. Final Presentation'
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Requirements Analysis
    doc.add_heading('1. Requirements Analysis', 1)
    
    doc.add_heading('1.1 Project Overview', 2)
    doc.add_paragraph(
        'The Service Request / Ticketing System is a web-based application designed to help organizations '
        'manage and track service requests efficiently. The system provides a user-friendly interface for '
        'creating, viewing, updating, and managing tickets throughout their lifecycle.'
    )
    
    doc.add_heading('1.2 Functional Requirements', 2)
    doc.add_paragraph('FR1: Users must be able to create new service request tickets with:')
    doc.add_paragraph('   - Ticket title (required)', style='List Bullet')
    doc.add_paragraph('   - Description (required)', style='List Bullet')
    doc.add_paragraph('   - Requester name (required)', style='List Bullet')
    doc.add_paragraph('   - Category (General, Technical, Billing, Support, Feature Request)', style='List Bullet')
    doc.add_paragraph('   - Priority level (Low, Medium, High)', style='List Bullet')
    
    doc.add_paragraph('FR2: The system must display all tickets in a visually appealing card-based grid layout')
    doc.add_paragraph('FR3: Users must be able to view detailed information about any ticket')
    doc.add_paragraph('FR4: Users must be able to update ticket status (Open, In Progress, Resolved, Closed)')
    doc.add_paragraph('FR5: Users must be able to add comments to tickets for collaboration')
    doc.add_paragraph('FR6: The system must provide filtering capabilities:')
    doc.add_paragraph('   - Filter by ticket status', style='List Bullet')
    doc.add_paragraph('   - Filter by priority level', style='List Bullet')
    doc.add_paragraph('FR7: The system must provide search functionality to find tickets by title, description, or requester name')
    doc.add_paragraph('FR8: The system must display real-time statistics including total tickets and counts by status')
    doc.add_paragraph('FR9: The system must automatically assign unique IDs to each ticket')
    doc.add_paragraph('FR10: The system must track creation and update timestamps for each ticket')
    
    doc.add_heading('1.3 Non-Functional Requirements', 2)
    nfr = [
        'NFR1: The user interface must be visually appealing with modern design principles',
        'NFR2: The system must be responsive and work on desktop, tablet, and mobile devices',
        'NFR3: The system must provide smooth animations and transitions for better user experience',
        'NFR4: The system must have fast response times for all operations',
        'NFR5: The system must be easy to use with intuitive navigation',
        'NFR6: The system must store data persistently (using JSON file storage)',
        'NFR7: The system must be built using modern web technologies'
    ]
    for req in nfr:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('1.4 User Stories', 2)
    stories = [
        'As a user, I want to create a new ticket so that I can report issues or request services',
        'As a user, I want to view all my tickets in one place so that I can track their status',
        'As a user, I want to filter tickets by status so that I can focus on specific types of tickets',
        'As a user, I want to search for tickets so that I can quickly find specific requests',
        'As a user, I want to see statistics about tickets so that I can understand the overall status',
        'As a user, I want to add comments to tickets so that I can provide additional information',
        'As a user, I want to update ticket status so that I can track progress'
    ]
    for story in stories:
        doc.add_paragraph(story, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. System Architecture Design
    doc.add_heading('2. System Architecture Design', 1)
    
    doc.add_heading('2.1 Architecture Overview', 2)
    doc.add_paragraph(
        'The system follows a client-server architecture with a RESTful API backend and a modern '
        'single-page application frontend. The architecture is designed to be simple, scalable, and maintainable.'
    )
    
    doc.add_heading('2.2 Technology Stack', 2)
    doc.add_paragraph('Backend Technologies:', style='Heading 3')
    backend = [
        'Python 3.x - Programming language',
        'Flask 3.0.0 - Web framework for building RESTful APIs',
        'JSON - Data storage format',
        'UUID - For generating unique ticket identifiers'
    ]
    for tech in backend:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_paragraph('Frontend Technologies:', style='Heading 3')
    frontend = [
        'HTML5 - Markup language',
        'CSS3 - Styling with modern features (gradients, animations, flexbox, grid)',
        'JavaScript (Vanilla) - Client-side interactivity',
        'Font Awesome 6.4.0 - Icon library',
        'Google Fonts (Inter) - Typography'
    ]
    for tech in frontend:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('2.3 System Components', 2)
    doc.add_paragraph('1. Backend Server (Flask Application)', style='Heading 3')
    doc.add_paragraph('The Flask application serves as the API server, handling all business logic and data operations.')
    
    doc.add_paragraph('2. Frontend Client (Web Interface)', style='Heading 3')
    doc.add_paragraph('The frontend is a single-page application that communicates with the backend via AJAX requests.')
    
    doc.add_paragraph('3. Data Storage Layer', style='Heading 3')
    doc.add_paragraph('Currently uses JSON file-based storage (tickets.json) for simplicity.')
    
    doc.add_heading('2.4 API Architecture', 2)
    doc.add_paragraph('The system implements a RESTful API with the following endpoints:')
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Endpoint'
    hdr[1].text = 'Description'
    
    endpoints = [
        ('GET /api/tickets', 'Retrieve all tickets, sorted by creation date'),
        ('POST /api/tickets', 'Create a new ticket'),
        ('GET /api/tickets/<id>', 'Retrieve a specific ticket by ID'),
        ('PUT /api/tickets/<id>', 'Update an existing ticket'),
        ('DELETE /api/tickets/<id>', 'Delete a ticket'),
        ('POST /api/tickets/<id>/comments', 'Add a comment to a ticket'),
        ('GET /api/stats', 'Get statistics about tickets')
    ]
    
    for endpoint, desc in endpoints:
        row = table.add_row().cells
        row[0].text = endpoint
        row[1].text = desc
    
    doc.add_heading('2.5 Data Model', 2)
    doc.add_paragraph('Ticket Data Structure:')
    add_code_block(doc, '''{
    "id": "unique-uuid",
    "title": "string",
    "description": "string",
    "priority": "low|medium|high",
    "status": "open|in_progress|resolved|closed",
    "category": "general|technical|billing|support|feature",
    "requester": "string",
    "created_at": "ISO 8601 timestamp",
    "updated_at": "ISO 8601 timestamp",
    "assigned_to": "string|null",
    "comments": [...]
}''')
    
    doc.add_heading('2.6 User Interface Design', 2)
    ui_features = [
        'Gradient backgrounds for visual appeal',
        'Glassmorphism effects with backdrop blur',
        'Card-based layout for ticket display',
        'Color-coded badges for status and priority',
        'Smooth animations and transitions',
        'Responsive grid layout',
        'Modal dialogs for ticket creation and details',
        'Real-time statistics dashboard'
    ]
    for feature in ui_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. MVP Implementation
    doc.add_heading('3. MVP (Minimum Viable Product) Implementation', 1)
    
    doc.add_heading('3.1 Implementation Overview', 2)
    doc.add_paragraph(
        'The MVP focuses on core functionality required for a functional ticketing system. All essential '
        'features have been implemented to provide a complete user experience.'
    )
    
    doc.add_heading('3.2 Core Features Implemented', 2)
    
    doc.add_paragraph('3.2.1 Ticket Creation', style='Heading 3')
    doc.add_paragraph('Users can create new tickets through a modal dialog with validation for required fields.')
    add_code_block(doc, '''@app.route('/api/tickets', methods=['POST'])
def create_ticket():
    data = request.json
    new_ticket = {
        'id': str(uuid4()),
        'title': data.get('title', ''),
        'description': data.get('description', ''),
        'priority': data.get('priority', 'medium'),
        'status': 'open',
        'category': data.get('category', 'general'),
        'requester': data.get('requester', 'Anonymous'),
        'created_at': datetime.now().isoformat(),
        'updated_at': datetime.now().isoformat(),
        'assigned_to': None,
        'comments': []
    }
    tickets.append(new_ticket)
    save_tickets(tickets)
    return jsonify(new_ticket), 201''')
    
    doc.add_paragraph('3.2.2 Ticket Display', style='Heading 3')
    doc.add_paragraph('Tickets are displayed in a responsive grid layout with cards showing key information.')
    
    doc.add_paragraph('3.2.3 Filtering and Search', style='Heading 3')
    doc.add_paragraph('The system provides real-time filtering by status and priority, plus text-based search.')
    
    doc.add_paragraph('3.2.4 Statistics Dashboard', style='Heading 3')
    doc.add_paragraph('Real-time statistics are displayed showing total tickets and counts by status.')
    
    doc.add_heading('3.3 File Structure', 2)
    add_code_block(doc, '''Ticketing System/
??? app.py                 # Flask backend
??? templates/
?   ??? index.html        # Main HTML template
??? static/
?   ??? style.css         # CSS styling
?   ??? script.js         # JavaScript
??? tickets.json          # Data storage
??? requirements.txt      # Dependencies
??? README.md            # Documentation''')
    
    doc.add_page_break()
    
    # 4. Testing
    doc.add_heading('4. Testing', 1)
    
    doc.add_heading('4.1 Testing Strategy', 2)
    doc.add_paragraph('Testing was performed manually throughout development to ensure all features work correctly.')
    
    doc.add_heading('4.2 Functional Testing', 2)
    
    doc.add_paragraph('4.2.1 Ticket Creation Testing', style='Heading 3')
    tests = [
        'Test creating a ticket with all required fields',
        'Test creating a ticket with missing required fields (validation)',
        'Test creating tickets with different priorities and categories',
        'Verify ticket appears in list after creation',
        'Verify ticket has unique ID and correct timestamps'
    ]
    for test in tests:
        doc.add_paragraph(test, style='List Bullet')
    
    doc.add_paragraph('4.2.2 Filtering and Search Testing', style='Heading 3')
    filter_tests = [
        'Test filtering by status (all, open, in_progress, resolved, closed)',
        'Test filtering by priority (all, high, medium, low)',
        'Test search functionality with various search terms',
        'Test combination of filters and search',
        'Verify results update in real-time'
    ]
    for test in filter_tests:
        doc.add_paragraph(test, style='List Bullet')
    
    doc.add_paragraph('4.2.3 Ticket Update Testing', style='Heading 3')
    update_tests = [
        'Test updating ticket status',
        'Verify updated timestamp changes',
        'Test adding comments to tickets',
        'Verify comments display correctly'
    ]
    for test in update_tests:
        doc.add_paragraph(test, style='List Bullet')
    
    doc.add_heading('4.3 User Interface Testing', 2)
    ui_tests = [
        'Test modal dialogs open and close correctly',
        'Test form validation and error messages',
        'Test responsive design on mobile devices',
        'Test hover effects and animations',
        'Test color coding for status and priority'
    ]
    for test in ui_tests:
        doc.add_paragraph(test, style='List Bullet')
    
    doc.add_heading('4.4 Test Results Summary', 2)
    doc.add_paragraph(
        'All core functionality has been tested and verified to work correctly. The system handles edge '
        'cases appropriately and provides a smooth user experience.'
    )
    
    doc.add_page_break()
    
    # 5. Documentation
    doc.add_heading('5. Documentation', 1)
    
    doc.add_heading('5.1 Code Documentation', 2)
    doc.add_paragraph('The codebase includes inline comments and docstrings explaining key functions.')
    
    doc.add_heading('5.2 User Documentation', 2)
    doc.add_paragraph('A comprehensive README.md file includes:')
    readme_items = [
        'Project overview and features',
        'Installation instructions',
        'Usage guide for all features',
        'API endpoint documentation',
        'Project structure explanation',
        'Customization guide'
    ]
    for item in readme_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.3 Installation Guide', 2)
    install_steps = [
        '1. Ensure Python 3.x is installed',
        '2. Navigate to the project directory',
        '3. Install dependencies: pip install -r requirements.txt',
        '4. Run the application: python app.py',
        '5. Open browser and navigate to http://localhost:5000'
    ]
    for step in install_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # 6. Final Presentation
    doc.add_heading('6. Final Presentation', 1)
    
    doc.add_heading('6.1 Project Summary', 2)
    doc.add_paragraph(
        'The Service Request / Ticketing System is a fully functional web application that successfully '
        'implements all MVP requirements.'
    )
    
    doc.add_heading('6.2 Key Achievements', 2)
    achievements = [
        'Successfully implemented all core functional requirements',
        'Created a visually stunning and modern user interface',
        'Built a responsive design that works on all devices',
        'Implemented real-time filtering and search functionality',
        'Developed a complete RESTful API backend',
        'Created comprehensive documentation'
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_heading('6.3 Technical Highlights', 2)
    highlights = [
        'RESTful API architecture for scalable backend',
        'Modern frontend with vanilla JavaScript',
        'JSON-based data storage for simplicity',
        'Responsive CSS Grid and Flexbox layouts',
        'Smooth animations and transitions',
        'Real-time statistics dashboard'
    ]
    for highlight in highlights:
        doc.add_paragraph(highlight, style='List Bullet')
    
    doc.add_heading('6.4 Future Enhancements', 2)
    future = [
        'User authentication and authorization',
        'Email notifications for ticket updates',
        'File attachment support',
        'Ticket assignment to team members',
        'Advanced reporting and analytics',
        'Database integration (PostgreSQL or SQLite)',
        'Real-time updates using WebSockets'
    ]
    for feature in future:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('6.5 Conclusion', 2)
    doc.add_paragraph(
        'The Service Request / Ticketing System successfully delivers a complete MVP with all core '
        'functionality. The system demonstrates modern web development practices, clean code architecture, '
        'and attention to user experience. The project is ready for deployment and can serve as a solid '
        'foundation for future enhancements.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('--- End of Document ---')
    
    # Save
    filename = 'Ticketing_System_Documentation.docx'
    doc.save(filename)
    print(f'Documentation created successfully: {filename}')
    return filename

if __name__ == '__main__':
    create_documentation()
