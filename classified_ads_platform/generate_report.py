"""
Script to generate a comprehensive Word document report for the Classified Ads Platform
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Classified Ads Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Requirements Analysis',
        '2. System Architecture Design',
        '3. MVP (Minimum Viable Product) Implementation',
        '4. Testing',
        '5. Documentation',
        '6. Final Presentation'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Requirements Analysis
    doc.add_heading('1. Requirements Analysis', 1)
    
    doc.add_heading('1.1 Functional Requirements', 2)
    doc.add_paragraph('The Classified Ads Platform must provide the following core functionalities:')
    
    func_req = [
        ('User Management', [
            'User registration with username, email, and password',
            'User authentication (login/logout)',
            'Password hashing for security',
            'User session management'
        ]),
        ('Ad Management', [
            'Create new classified ads with title, description, price, location, and contact information',
            'Edit existing ads (only by the ad owner)',
            'Delete ads (only by the ad owner)',
            'View detailed ad information',
            'List all active ads with pagination'
        ]),
        ('Category System', [
            'Organize ads into predefined categories',
            'Browse ads by category',
            'Default categories: Electronics, Vehicles, Real Estate, Furniture, Clothing, Books, Sports, Services, Other'
        ]),
        ('Search and Filter', [
            'Search ads by keywords (title and description)',
            'Filter ads by category',
            'Filter ads by price range (min/max)',
            'Filter ads by location'
        ]),
        ('User Dashboard', [
            'View all ads posted by the logged-in user',
            'Quick access to edit/delete own ads'
        ])
    ]
    
    for req_name, req_items in func_req:
        doc.add_paragraph(req_name, style='Heading 3')
        for item in req_items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('1.2 Non-Functional Requirements', 2)
    nfr_items = [
        'Security: Passwords must be hashed using Werkzeug security functions',
        'Usability: Responsive web interface that works on desktop and mobile devices',
        'Performance: Efficient database queries with pagination for large datasets',
        'Maintainability: Clean code structure following Flask best practices',
        'Scalability: SQLite database (can be upgraded to PostgreSQL/MySQL for production)',
        'User Experience: Intuitive navigation and clear visual feedback'
    ]
    for item in nfr_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('1.3 Technology Stack', 2)
    tech_stack = [
        ('Backend Framework', 'Flask 3.0.0 - Lightweight Python web framework'),
        ('Database', 'SQLite with SQLAlchemy ORM for database operations'),
        ('Authentication', 'Flask-Login for session management and user authentication'),
        ('Forms', 'Flask-WTF and WTForms for form handling and validation'),
        ('Frontend', 'HTML5, CSS3 with responsive design'),
        ('Security', 'Werkzeug for password hashing and CSRF protection')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    for tech, desc in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = desc
    
    doc.add_page_break()
    
    # 2. System Architecture Design
    doc.add_heading('2. System Architecture Design', 1)
    
    doc.add_heading('2.1 Architecture Overview', 2)
    doc.add_paragraph(
        'The application follows a Model-View-Controller (MVC) architectural pattern, '
        'adapted for Flask\'s structure. The system is organized into distinct layers '
        'for maintainability and scalability.'
    )
    
    doc.add_heading('2.2 System Components', 2)
    
    components = [
        ('Models Layer (models.py)', [
            'User Model: Stores user account information (username, email, password hash)',
            'Category Model: Defines ad categories with name and description',
            'Ad Model: Contains ad details (title, description, price, location, contact info, timestamps)',
            'Relationships: User has many Ads, Category has many Ads, Ad belongs to User and Category'
        ]),
        ('Forms Layer (forms.py)', [
            'RegistrationForm: User registration with validation',
            'LoginForm: User authentication',
            'AdForm: Create and edit ads with field validation',
            'SearchForm: Search and filter functionality'
        ]),
        ('Controller Layer (app.py)', [
            'Route handlers for all application endpoints',
            'Business logic and request processing',
            'Database initialization and default data seeding',
            'Authentication and authorization checks'
        ]),
        ('View Layer (templates/)', [
            'Base template with navigation and layout',
            'Home page with category browsing and recent ads',
            'User authentication pages (login/register)',
            'Ad management pages (create, edit, detail, list)',
            'Search page with filtering options'
        ]),
        ('Static Assets (static/)', [
            'CSS stylesheet with responsive design',
            'Modern UI with consistent color scheme and typography'
        ])
    ]
    
    for comp_name, comp_items in components:
        doc.add_paragraph(comp_name, style='Heading 3')
        for item in comp_items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('2.3 Database Schema', 2)
    doc.add_paragraph('The database consists of three main tables:')
    
    schema_desc = [
        ('User Table', [
            'id (Primary Key)',
            'username (Unique)',
            'email (Unique)',
            'password_hash',
            'created_at'
        ]),
        ('Category Table', [
            'id (Primary Key)',
            'name (Unique)',
            'description'
        ]),
        ('Ad Table', [
            'id (Primary Key)',
            'title',
            'description',
            'price',
            'location',
            'contact_info',
            'created_at',
            'updated_at',
            'is_active',
            'user_id (Foreign Key ? User)',
            'category_id (Foreign Key ? Category)'
        ])
    ]
    
    for table_name, fields in schema_desc:
        doc.add_paragraph(table_name, style='Heading 3')
        for field in fields:
            doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('2.4 Request Flow', 2)
    doc.add_paragraph(
        '1. User makes HTTP request to Flask application\n'
        '2. Flask routes request to appropriate handler function\n'
        '3. Handler validates input using forms and checks authentication\n'
        '4. Business logic executes (database queries, data processing)\n'
        '5. Response rendered using Jinja2 templates\n'
        '6. HTML response sent to user\'s browser'
    )
    
    doc.add_page_break()
    
    # 3. MVP Implementation
    doc.add_heading('3. MVP (Minimum Viable Product) Implementation', 1)
    
    doc.add_heading('3.1 Core Features Implemented', 2)
    
    mvp_features = [
        ('User Authentication System', [
            'Complete registration and login functionality',
            'Secure password hashing using Werkzeug',
            'Session management with Flask-Login',
            'Protected routes requiring authentication'
        ]),
        ('Ad Management System', [
            'Create ads with all required fields',
            'Edit ads (with ownership verification)',
            'Delete ads (with ownership verification)',
            'View individual ad details',
            'List all active ads with pagination'
        ]),
        ('Category Organization', [
            '9 predefined categories',
            'Category-based browsing',
            'Category filtering in search'
        ]),
        ('Search and Filter', [
            'Full-text search in titles and descriptions',
            'Category filtering',
            'Price range filtering',
            'Location-based filtering',
            'Combined filter support'
        ]),
        ('User Dashboard', [
            'View all user\'s own ads',
            'Quick access to edit/delete functionality'
        ]),
        ('Responsive Web Interface', [
            'Modern, clean design',
            'Mobile-responsive layout',
            'Intuitive navigation',
            'User-friendly forms with validation feedback'
        ])
    ]
    
    for feature_name, feature_items in mvp_features:
        doc.add_paragraph(feature_name, style='Heading 3')
        for item in feature_items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.2 File Structure', 2)
    doc.add_paragraph('The project follows a well-organized structure:')
    
    file_structure = [
        'app.py - Main Flask application with all routes',
        'models.py - Database models (User, Category, Ad)',
        'forms.py - WTForms for form validation',
        'requirements.txt - Python dependencies',
        'README.md - Project documentation',
        'templates/ - HTML templates (8 files)',
        'static/style.css - Stylesheet',
        '.gitignore - Git ignore rules'
    ]
    
    for item in file_structure:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.3 Key Implementation Details', 2)
    
    impl_details = [
        ('Database Initialization', 
         'Database tables are automatically created on application startup. '
         'Default categories are seeded if they don\'t exist.'),
        ('Security Measures',
         'Passwords are hashed using Werkzeug\'s generate_password_hash. '
         'CSRF protection enabled via Flask-WTF. User sessions managed securely.'),
        ('Form Validation',
         'Client-side and server-side validation using WTForms validators. '
         'Custom validators for username/email uniqueness checks.'),
        ('Error Handling',
         '404 errors for missing resources. Flash messages for user feedback. '
         'Graceful handling of database errors.'),
        ('Pagination',
         'Efficient pagination for ad listings using Flask-SQLAlchemy paginate method. '
         '12 ads per page to optimize performance.')
    ]
    
    for detail_name, detail_desc in impl_details:
        doc.add_paragraph(detail_name, style='Heading 3')
        doc.add_paragraph(detail_desc)
    
    doc.add_page_break()
    
    # 4. Testing
    doc.add_heading('4. Testing', 1)
    
    doc.add_heading('4.1 Testing Strategy', 2)
    doc.add_paragraph(
        'The application was tested through manual testing and functional verification. '
        'All core features were validated to ensure they work as expected.'
    )
    
    doc.add_heading('4.2 Test Cases', 2)
    
    test_cases = [
        ('User Registration', [
            '? Valid registration creates new user account',
            '? Duplicate username/email rejected',
            '? Password validation enforced',
            '? Successful registration redirects to login'
        ]),
        ('User Authentication', [
            '? Valid credentials allow login',
            '? Invalid credentials show error message',
            '? Logged-in users redirected from login/register pages',
            '? Logout successfully ends session'
        ]),
        ('Ad Creation', [
            '? Authenticated users can create ads',
            '? All required fields validated',
            '? Ad successfully saved to database',
            '? Redirect to ad detail page after creation'
        ]),
        ('Ad Management', [
            '? Users can edit their own ads',
            '? Users cannot edit others\' ads',
            '? Users can delete their own ads',
            '? Users cannot delete others\' ads',
            '? Ad updates reflect immediately'
        ]),
        ('Search and Filter', [
            '? Keyword search finds matching ads',
            '? Category filter works correctly',
            '? Price range filter functions properly',
            '? Location filter operates as expected',
            '? Combined filters work together'
        ]),
        ('Database Operations', [
            '? Database tables created on startup',
            '? Default categories seeded correctly',
            '? Foreign key relationships maintained',
            '? Data persistence verified'
        ]),
        ('User Interface', [
            '? All pages render correctly',
            '? Navigation works on all pages',
            '? Forms display validation errors',
            '? Flash messages appear appropriately',
            '? Responsive design works on mobile'
        ])
    ]
    
    for test_name, test_items in test_cases:
        doc.add_paragraph(test_name, style='Heading 3')
        for item in test_items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.3 Known Issues and Resolutions', 2)
    doc.add_paragraph('Issue: Database tables not created on first run')
    doc.add_paragraph('Resolution: Added database initialization on app startup with app context', style='List Bullet 2')
    doc.add_paragraph()
    doc.add_paragraph('All identified issues have been resolved. The application runs successfully.')
    
    doc.add_page_break()
    
    # 5. Documentation
    doc.add_heading('5. Documentation', 1)
    
    doc.add_heading('5.1 Code Documentation', 2)
    doc.add_paragraph(
        'The codebase includes comprehensive documentation:'
    )
    doc_items = [
        'Function docstrings for all route handlers',
        'Class docstrings for all models',
        'Inline comments for complex logic',
        'Clear variable and function naming conventions'
    ]
    for item in doc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.2 User Documentation', 2)
    doc.add_paragraph('README.md file includes:')
    readme_items = [
        'Project overview and features',
        'Installation instructions',
        'Usage guide',
        'Project structure explanation'
    ]
    for item in readme_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.3 API Documentation', 2)
    doc.add_paragraph('Available Routes:')
    
    routes = [
        ('GET /', 'Home page - displays recent ads and categories'),
        ('GET /register', 'Registration page'),
        ('POST /register', 'Process registration'),
        ('GET /login', 'Login page'),
        ('POST /login', 'Process login'),
        ('GET /logout', 'Logout user'),
        ('GET /create_ad', 'Create ad form (requires login)'),
        ('POST /create_ad', 'Process ad creation'),
        ('GET /ad/<id>', 'View ad details'),
        ('GET /edit_ad/<id>', 'Edit ad form (requires login, owner only)'),
        ('POST /edit_ad/<id>', 'Process ad update'),
        ('POST /delete_ad/<id>', 'Delete ad (requires login, owner only)'),
        ('GET /my_ads', 'User\'s ads dashboard (requires login)'),
        ('GET /search', 'Search page'),
        ('POST /search', 'Process search query'),
        ('GET /category/<id>', 'View ads by category')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Route'
    hdr_cells[1].text = 'Description'
    for route, desc in routes:
        row_cells = table.add_row().cells
        row_cells[0].text = route
        row_cells[1].text = desc
    
    doc.add_page_break()
    
    # 6. Final Presentation
    doc.add_heading('6. Final Presentation', 1)
    
    doc.add_heading('6.1 Project Summary', 2)
    doc.add_paragraph(
        'The Classified Ads Platform is a fully functional web application that allows users to '
        'post, browse, search, and manage classified advertisements. The platform successfully '
        'implements all MVP features including user authentication, ad management, category '
        'organization, and advanced search capabilities.'
    )
    
    doc.add_heading('6.2 Achievements', 2)
    achievements = [
        'Complete user authentication system with secure password handling',
        'Full CRUD operations for classified ads',
        'Comprehensive search and filtering system',
        'Responsive web interface with modern design',
        'Well-structured, maintainable codebase',
        'Comprehensive documentation',
        'Production-ready database structure'
    ]
    for item in achievements:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.3 Technology Highlights', 2)
    doc.add_paragraph(
        'The project demonstrates proficiency in:'
    )
    tech_highlights = [
        'Flask web framework and routing',
        'SQLAlchemy ORM for database operations',
        'User authentication and session management',
        'Form validation and CSRF protection',
        'Jinja2 templating',
        'Responsive CSS design',
        'RESTful API design principles'
    ]
    for item in tech_highlights:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.4 Future Enhancements', 2)
    doc.add_paragraph('Potential improvements for future versions:')
    future_items = [
        'Image upload and management for ads',
        'User profiles with avatars',
        'Messaging system between buyers and sellers',
        'Email notifications for new ads in categories',
        'Advanced search with sorting options',
        'Admin panel for category management',
        'Ad favorites/watchlist functionality',
        'Rating and review system',
        'Payment integration for premium listings',
        'Migration to PostgreSQL for better scalability'
    ]
    for item in future_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.5 Conclusion', 2)
    doc.add_paragraph(
        'The Classified Ads Platform successfully delivers a complete MVP with all core '
        'functionalities working as expected. The application is well-structured, secure, '
        'and ready for deployment. The codebase follows best practices and is maintainable '
        'for future development.'
    )
    
    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph('--- End of Report ---')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer.runs[0]
    footer_format.italic = True
    footer_format.font.size = Pt(10)
    
    # Save document
    doc.save('Classified_Ads_Platform_Report.docx')
    print("Report generated successfully: Classified_Ads_Platform_Report.docx")

if __name__ == '__main__':
    create_report()
